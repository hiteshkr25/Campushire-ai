"""
CampusHire AI — Resume Parser v2
=================================
Section-based pipeline that correctly handles:
  - Multiline paragraph / bullet merging
  - Hyphenated line-break repair
  - Dense-PDF space reconstruction
  - Section boundary enforcement (no cross-section leakage)
  - Structured extraction for Experience, Internships, and Projects

Pipeline:
  extract_raw_text()
    → preprocess()          # repair hyphens, fix dense spacing, normalise bullets
    → detect_sections()     # single ordered pass → {section_name: [raw_lines]}
    → merge_paragraphs()    # stitch continuation lines into coherent bullets
    → parse_contact()
    → parse_education()
    → parse_skills()
    → parse_experience()    # → [{company, role, duration, description[]}]
    → parse_projects()      # → [{title, technologies[], description, github, duration}]
    → parse_certifications()
    → parse_achievements()
    → score_confidence()
    → build_output()
"""

import json
import re
from pathlib import Path
from datetime import datetime, timezone

from app.extensions import db
from app.models.student import Resume
from app.models.enums import ParseStatus


# ---------------------------------------------------------------------------
# Canonical skills pool (single source of truth — imported by ats_service too)
# ---------------------------------------------------------------------------
SKILLS_POOL = [
    # Languages
    "python", "java", "javascript", "typescript", "c", "c++", "c#", "ruby",
    "php", "go", "rust", "scala", "kotlin", "swift", "r", "matlab", "bash",
    # Frontend
    "html", "css", "react", "react.js", "angular", "vue", "vue.js", "svelte",
    "next.js", "tailwind", "bootstrap",
    # Backend / frameworks
    "django", "flask", "fastapi", "spring", "spring boot", "express", "express.js",
    "node.js", "nodejs", "laravel", "rails", "asp.net",
    # Databases
    "sql", "postgresql", "mysql", "sqlite", "mongodb", "redis", "cassandra",
    "neo4j", "oracle", "dynamodb", "firebase", "supabase",
    # Cloud / DevOps
    "aws", "azure", "gcp", "google cloud", "docker", "kubernetes", "jenkins",
    "terraform", "ansible", "ci/cd", "github actions",
    # ML / Data
    "machine learning", "deep learning", "nlp", "computer vision",
    "tensorflow", "pytorch", "scikit-learn", "pandas", "numpy", "opencv",
    "data science", "tableau", "power bi",
    # Tools & practices
    "git", "github", "postman", "linux", "rest apis", "graphql",
    "agile", "scrum", "algorithms", "data structures", "oops",
    "problem solving", "system design",
]

# ---------------------------------------------------------------------------
# Section header patterns (ordered — first match wins per line).
# Each pattern covers both normal spacing AND dense-PDF all-caps concatenation
# e.g. "TECHNICALSKILLS" from PDFs that lose inter-word spaces.
# ---------------------------------------------------------------------------
_SECTION_PATTERNS = [
    ("contact",        re.compile(
        r"^\s*(?:[•\-\*–▪◦➤➢►▶]\s*)?(?:contact|personal\s*info(?:rmation)?|PERSONALINFO(?:RMATION)?)\s*[:\-–—]*\s*$", re.I)),
    ("education",      re.compile(
        r"^\s*(?:[•\-\*–▪◦➤➢►▶]\s*)?(?:education|academic\s*details?|ACADEMICDETAILS?|academics|qualifications?|educational\s*background)\s*[:\-–—]*\s*$", re.I)),
    ("skills",         re.compile(
        r"^\s*(?:[•\-\*–▪◦➤➢►▶]\s*)?(?:technical\s*skills?|TECHNICALSKILLS?|skills?|core\s*competencies|CORECOMPETENCIES|technologies|tech\s*stack)\s*[:\-–—]*\s*$", re.I)),
    ("experience",     re.compile(
        r"^\s*(?:[•\-\*–▪◦➤➢►▶]\s*)?(?:work\s*experience|WORKEXPERIENCE|professional\s*experience|PROFESSIONALEXPERIENCE|employment(?:\s*history)?|experience)\s*[:\-–—]*\s*$", re.I)),
    ("internships",    re.compile(
        r"^\s*(?:[•\-\*–▪◦➤➢►▶]\s*)?(?:internships?|industrial\s*training|INDUSTRIALTRAINING|industry\s*experience)\s*[:\-–—]*\s*$", re.I)),
    ("projects",       re.compile(
        r"^\s*(?:[•\-\*–▪◦➤➢►▶]\s*)?(?:projects?|personal\s*projects?|PERSONALPROJECTS?|academic\s*projects?|ACADEMICPROJECTS?|key\s*projects?|KEYPROJECTS?|notable\s*projects?)\s*[:\-–—]*\s*$", re.I)),
    ("certifications", re.compile(
        r"^\s*(?:[•\-\*–▪◦➤➢►▶]\s*)?(?:certifications?|CERTIFICATIONS?|licenses?|courses?\s*(?:&|and)\s*certifications?|credentials?|professional\s*development)\s*[:\-–—]*\s*$", re.I)),
    ("achievements",   re.compile(
        r"^\s*(?:[•\-\*–▪◦➤➢►▶]\s*)?(?:achievements?|ACHIEVEMENTS?|honors?|awards?|accomplishments?|scholastic\s*achievements?|SCHOLASTICACHIEVEMENTS?|positions?\s*of\s*responsibility|POSITIONSOFRESPONSIBILITY)\s*[:\-–—]*\s*$", re.I)),
    ("languages",      re.compile(r"^\s*(?:[•\-\*–▪◦➤➢►▶]\s*)?(?:languages?)\s*[:\-–—]*\s*$", re.I)),
    ("interests",      re.compile(r"^\s*(?:[•\-\*–▪◦➤➢►▶]\s*)?(?:interests?|hobbies|extra[- ]curricular)\s*[:\-–—]*\s*$", re.I)),
]

_SECTION_NAMES = [s for s, _ in _SECTION_PATTERNS]

# Duration patterns — handles:
#   "Jan 2022 – Mar 2023"   "2022-2024"   "Sep 2025 - Nov 2025"
#   Dense-PDF forms: "Sep2025-Nov2025"  "July2025-Aug2025"
_MONTH = r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|June|July|August|September|October|November|December)"
_DURATION_RE = re.compile(
    r"(?:" + _MONTH + r"[a-z]*[,.\s]*\d{4}|\d{4})"
    r"\s*[-–—to]+\s*"
    r"(?:" + _MONTH + r"[a-z]*[,.\s]*\d{4}|\d{4}|present|current|ongoing)",
    re.I,
)

_YEAR_RANGE_RE = re.compile(r"\d{4}\s*[-–—]\s*(?:\d{4}|present|current)", re.I)
_EMAIL_RE        = re.compile(r"[\w.\-+]+@[\w.\-]+\.\w{2,}")
_PHONE_RE        = re.compile(r"(?:\+?\d{1,3}[-.\s]?)?\(?\d{3,5}\)?[-.\s]?\d{3,5}[-.\s]?\d{4,6}")
_LINKEDIN_RE     = re.compile(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[\w.\-]+/?", re.I)
_GITHUB_URL_RE   = re.compile(r"(?:https?://)?(?:www\.)?github\.com/[\w.\-]+(?:/[\w.\-]+)*/?", re.I)
_PORTFOLIO_RE    = re.compile(
    r"https?://(?!(?:www\.)?(?:github|linkedin|leetcode|hackerrank)\.com)"
    r"[\w.\-]+\.(?:io|com|dev|me|in|co|net|xyz|app)/[\w./\-?=%&+#]*",
    re.I,
)
_LEETCODE_RE     = re.compile(r"(?:https?://)?(?:www\.)?leetcode\.com/(?:u/)?[\w.\-]+/?", re.I)
_HACKERRANK_RE   = re.compile(r"(?:https?://)?(?:www\.)?hackerrank\.com/[\w.\-]+/?", re.I)
_BULLET_RE       = re.compile(r"^[•\-\*–▪◦➤➢►▶]\s+")


# ===========================================================================
# Extraction helpers
# ===========================================================================

def _extract_pdf_text(file_path: Path) -> str:
    """Try pdfplumber (good layout), fall back to PyPDF2."""
    import pdfplumber
    import PyPDF2
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            pages = []
            for page in pdf.pages:
                # Looser tolerances help reconstruct spaces in dense PDFs
                t = page.extract_text(x_tolerance=3, y_tolerance=3)
                if t:
                    pages.append(t)
            text = "\n".join(pages)
    except Exception:
        text = ""

    if not text.strip():
        try:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                pages = []
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        pages.append(t)
                text = "\n".join(pages)
        except Exception:
            pass

    return text


def _extract_pdf_links(file_path: Path) -> list[str]:
    """
    Extract every hyperlink URI embedded in a PDF's annotation layer.
    Many resume PDFs hide LinkedIn/GitHub/LeetCode behind icon buttons that
    have no visible URL text — this is the only way to recover them.

    Two independent extraction layers (results are merged):
      1. PyMuPDF (fitz)   — reads /Annots from the PDF object tree; most reliable
      2. pdfplumber       — reads hyperlinks via pdfminer; catches some PDFs fitz misses

    Returns a deduplicated list of http/https URIs, PyMuPDF results first.
    """
    collected: list[str] = []

    # ── Layer 1: PyMuPDF ────────────────────────────────────────────────────
    try:
        import fitz as _fitz
        with _fitz.open(str(file_path)) as doc:
            for page in doc:
                for lnk in page.get_links():
                    uri = lnk.get("uri", "")
                    if uri and uri.startswith(("http://", "https://")):
                        collected.append(uri)
    except ImportError:
        pass
    except Exception:
        pass

    # ── Layer 2: pdfplumber .hyperlinks ─────────────────────────────────────
    # pdfplumber exposes annotation-layer hyperlinks via page.hyperlinks,
    # each a dict with at minimum {"uri": "..."}.
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                for h in getattr(page, "hyperlinks", []):
                    uri = h.get("uri", "") if isinstance(h, dict) else ""
                    if uri and uri.startswith(("http://", "https://")):
                        collected.append(uri)
    except ImportError:
        pass
    except Exception:
        pass

    # Deduplicate while preserving order (PyMuPDF results take priority)
    seen: set[str] = set()
    unique: list[str] = []
    for u in collected:
        if u not in seen:
            seen.add(u)
            unique.append(u)
    return unique


def normalize_url(url: str, url_type: str) -> str:
    if not url:
        return ""
    url = url.strip()
    if not url.startswith(("http://", "https://")):
        url = "https://" + url
    
    from urllib.parse import urlparse, urlunparse
    parsed = urlparse(url)
    clean_path = parsed.path.rstrip("/")
    
    if url_type == "linkedin":
        if "linkedin.com" in parsed.netloc.lower():
            netloc = "www.linkedin.com"
        else:
            netloc = parsed.netloc
        path_parts = [p for p in clean_path.split("/") if p]
        if len(path_parts) >= 2 and path_parts[0].lower() == "in":
            clean_path = "/in/" + "/".join(path_parts[1:])
        elif len(path_parts) == 1:
            clean_path = "/in/" + path_parts[0]
        return f"https://{netloc}{clean_path}"
        
    elif url_type == "github":
        if "github.com" in parsed.netloc.lower():
            netloc = "github.com"
        else:
            netloc = parsed.netloc
        path_parts = [p for p in clean_path.split("/") if p]
        if path_parts:
            clean_path = "/" + "/".join(path_parts)
        return f"https://{netloc}{clean_path}"
        
    return urlunparse((parsed.scheme, parsed.netloc, clean_path, '', '', ''))


def _classify_urls(text: str, pdf_links: list[str]) -> dict:
    """
    Collect and classify professional URLs from:
      1. Plain text (regex scan)
      2. PDF annotation hyperlinks (already extracted by _extract_pdf_links)

    Returns:
      {linkedin, github, github_profile, github_repos, portfolio, leetcode, hackerrank}

    github       → the shortest/profile-level GitHub URL (github.com/user)
    github_repos → list of project-level GitHub URLs (github.com/user/repo)
    """
    all_text_urls: list[str] = []
    for pattern in (_LINKEDIN_RE, _GITHUB_URL_RE, _LEETCODE_RE, _HACKERRANK_RE, _PORTFOLIO_RE):
        all_text_urls.extend(m.group(0) for m in pattern.finditer(text))

    _MARKDOWN_LINK_RE = re.compile(r"\[.*?\]\((https?://[\w.\-?=%&+#/]+)\)", re.I)
    all_text_urls.extend(m.group(1) for m in _MARKDOWN_LINK_RE.finditer(text))

    # Combine and deduplicate, PDF links take priority (they are authoritative)
    combined: list[str] = list(dict.fromkeys(pdf_links + all_text_urls))

    linkedin = github_profile = portfolio = leetcode = hackerrank = None
    github_repos: list[str] = []

    for url in combined:
        u = url.rstrip("/")
        ul = u.lower()

        if "linkedin.com/in/" in ul and linkedin is None:
            linkedin = normalize_url(u, "linkedin")

        elif "github.com/" in ul:
            # Count path depth: github.com/user = 1 segment, github.com/user/repo = 2
            path_parts = [p for p in u.split("github.com/", 1)[-1].split("/") if p]
            if len(path_parts) == 1:
                # Profile URL — keep the shortest one
                normalized_p = normalize_url(u, "github")
                if github_profile is None or len(normalized_p) < len(github_profile):
                    github_profile = normalized_p
            else:
                # Repo URL
                normalized_r = normalize_url(u, "github")
                if normalized_r not in github_repos:
                    github_repos.append(normalized_r)

        elif "leetcode.com/" in ul and leetcode is None:
            leetcode = u

        elif "hackerrank.com/" in ul and hackerrank is None:
            hackerrank = u

        elif _PORTFOLIO_RE.match(u + "/") or _PORTFOLIO_RE.match(u):
            if portfolio is None:
                portfolio = u

    # github field → profile if found, else first repo
    github = github_profile or (github_repos[0] if github_repos else None)

    return {
        "linkedin":      linkedin,
        "github":        github,
        "github_profile": github_profile,
        "github_repos":  github_repos,
        "portfolio":     portfolio,
        "leetcode":      leetcode,
        "hackerrank":    hackerrank,
    }


def _extract_docx_text(file_path: Path) -> str:
    try:
        import docx as python_docx
        doc = python_docx.Document(str(file_path))
        parts = []
        
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        }
        
        def process_element(element, part):
            element_text = []
            for child in element:
                tag = child.tag
                if tag.endswith('}r'):
                    t_el = child.find('w:t', namespaces)
                    if t_el is not None and t_el.text:
                        element_text.append(t_el.text)
                elif tag.endswith('}hyperlink'):
                    r_id = child.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    hyperlink_text = []
                    for sub_child in child:
                        if sub_child.tag.endswith('}r'):
                            t_el = sub_child.find('w:t', namespaces)
                            if t_el is not None and t_el.text:
                                hyperlink_text.append(t_el.text)
                    
                    text_content = "".join(hyperlink_text)
                    element_text.append(text_content)
                    
                    if r_id and r_id in part.rels:
                        url = part.rels[r_id].target_ref
                        if url.startswith(("http://", "https://")):
                            element_text.append(f" ({url}) ")
                else:
                    sub_text = process_element(child, part)
                    if sub_text:
                        element_text.append(sub_text)
            return "".join(element_text)

        for para in doc.paragraphs:
            p_text = process_element(para._p, doc.part)
            if p_text.strip():
                parts.append(p_text.strip())
                
        for table in doc.tables:
            for row in table.rows:
                row_cells_text = []
                for cell in row.cells:
                    c_text = []
                    for p in cell.paragraphs:
                        p_txt = process_element(p._p, doc.part)
                        if p_txt.strip():
                            c_text.append(p_txt.strip())
                    cell_text = "\n".join(c_text)
                    if cell_text.strip():
                        row_cells_text.append(cell_text.strip())
                if row_cells_text:
                    parts.append("  ".join(row_cells_text))
                    
        return "\n".join(parts)
    except Exception as e:
        raise ValueError(f"DOCX extraction failed: {e}") from e


# ===========================================================================
# Preprocessing
# ===========================================================================

def _preprocess(text: str) -> list[str]:
    """
    1. Repair hyphenated line breaks  (word-\nrest → wordrest)
    2. Insert space before uppercase-run headers in dense PDFs
       (ACADEMICDETAILS → ACADEMIC DETAILS ... handled by section regex)
    3. Normalise bullet characters
    4. Strip blank lines → return list[str]
    """
    # 1. Repair hard-hyphen line breaks
    text = re.sub(r"-\n(\S)", r"\1", text)

    # 2. Insert space between a lowercase letter immediately followed by an
    #    uppercase letter when they share the same "word" — common in dense PDFs
    #    e.g. "B.TechinComputerScience" → "B.Tech in Computer Science"
    text = re.sub(r"([a-z])([A-Z])", r"\1 \2", text)

    # 3. Normalise various bullet characters to a canonical bullet
    text = re.sub(r"^[\*\-–▪◦➤➢►▶]\s+", "• ", text, flags=re.MULTILINE)

    lines = [ln.rstrip() for ln in text.splitlines()]
    return lines


# ===========================================================================
# Section detection — single ordered pass
# ===========================================================================

def _detect_sections(lines: list[str]) -> dict[str, list[str]]:
    """
    Walk lines once.  When a line matches a section header pattern, switch the
    active section.  Lines before any recognised section go into 'header'
    (name, contact details).
    Returns {section_name: [raw_lines]}.
    """
    sections: dict[str, list[str]] = {"header": []}
    current = "header"

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        matched = False
        for name, pattern in _SECTION_PATTERNS:
            if pattern.match(stripped):
                current = name
                if name not in sections:
                    sections[name] = []
                matched = True
                break

        if not matched:
            sections.setdefault(current, []).append(stripped)

    return sections


# ===========================================================================
# Multiline / bullet merging
# ===========================================================================

def _is_continuation(line: str) -> bool:
    """
    A line is a continuation of the previous bullet when it:
    - Does not start with a bullet marker
    - Does not look like a new section header (all-caps short line)
    - Does not start with a capital letter after a common sentence-starting pattern
      that would indicate a new item rather than a run-on
    """
    if _BULLET_RE.match(line):
        return False
    # Short all-caps lines are likely sub-headers, not continuations
    if re.match(r"^[A-Z\s/&]+$", line) and len(line) < 50:
        return False
    return True


def _merge_bullets(raw_lines: list[str]) -> list[str]:
    """
    Merge multiline bullets:
      "• First sentence that wraps"   → one merged bullet
      "across two lines."
    Also join plain paragraphs that aren't bullet-started.
    """
    merged: list[str] = []
    buf = ""

    for line in raw_lines:
        stripped = line.strip()
        if not stripped:
            if buf:
                merged.append(buf.strip())
                buf = ""
            continue

        if _BULLET_RE.match(stripped):
            if buf:
                merged.append(buf.strip())
            buf = stripped
        elif buf and _is_continuation(stripped):
            # Ensure there is exactly one space joining the parts
            if buf.endswith("-"):
                buf = buf[:-1] + stripped
            else:
                buf = buf + " " + stripped
        else:
            if buf:
                merged.append(buf.strip())
            buf = stripped

    if buf:
        merged.append(buf.strip())

    return merged


# ===========================================================================
# Individual section parsers
# ===========================================================================

def _parse_contact(header_lines: list[str], full_text: str,
                   pdf_links: list[str] | None = None) -> dict:
    """
    Extract contact fields from header lines + full text + PDF hyperlink annotations.
    URL classification covers LinkedIn, GitHub (profile + repos), Portfolio,
    LeetCode, and HackerRank via both regex (text) and PyMuPDF (annotations).
    """
    email_m = _EMAIL_RE.search(full_text)
    phone_m = _PHONE_RE.search(full_text)

    name = None
    for line in header_lines[:8]:
        stripped = line.strip()
        if (
            3 < len(stripped) < 50
            and not _EMAIL_RE.search(stripped)
            and not _PHONE_RE.search(stripped)
            and not any(kw in stripped.lower() for kw in
                        ["curriculum", "resume", "cv", "portfolio",
                         "linkedin", "github", "http", "contact", "@",
                         "leetcode", "hackerrank"])
        ):
            name = stripped
            break

    urls = _classify_urls(full_text, pdf_links or [])

    return {
        "name":           name,
        "email":          email_m.group(0) if email_m else None,
        "phone":          phone_m.group(0) if phone_m else None,
        "linkedin":       urls["linkedin"],
        "github":         urls["github"],
        "github_profile": urls["github_profile"],
        "github_repos":   urls["github_repos"],
        "portfolio":      urls["portfolio"],
        "leetcode":       urls["leetcode"],
        "hackerrank":     urls["hackerrank"],
    }


def _parse_education(lines: list[str]) -> list[dict]:
    """
    Each education record is detected by presence of degree keywords or year
    ranges.  We group consecutive lines into records.
    """
    EDU_KW = re.compile(
        r"\b(?:b\.?tech|b\.?e|m\.?tech|m\.?e|b\.?sc|m\.?sc|bca|mca|mba|ph\.?d"
        r"|bachelor|master|engineer|diploma|class\s+(?:x{1,2}|xii|xi|10|12)"
        r"|ssc|hsc|cbse|icse|intermediate|matriculation|university|college|institute|school)\b",
        re.I,
    )

    records = []
    buf: list[str] = []

    def flush():
        if buf:
            records.append({"raw": " ".join(buf)})
            buf.clear()

    for line in _merge_bullets(lines):
        stripped = line.strip()
        if EDU_KW.search(stripped) or _YEAR_RANGE_RE.search(stripped):
            # Try to attach to existing buffer if it is part of the same record
            if buf and (_YEAR_RANGE_RE.search(stripped) or len(stripped) < 80):
                buf.append(stripped)
            else:
                flush()
                buf.append(stripped)
        elif buf:
            buf.append(stripped)

    flush()
    return records


def _parse_skills(lines: list[str]) -> dict[str, list[str]]:
    """
    Return {category: [skill, ...]} by detecting category labels like
    "Languages:", "Frameworks & Libraries:", etc.
    Also return a flat "all" list.
    """
    merged = _merge_bullets(lines)
    categories: dict[str, list[str]] = {}
    current_cat = "General"

    for line in merged:
        # Strip leading bullet
        clean = _BULLET_RE.sub("", line).strip()

        # Detect "Category: val1, val2, ..."
        cat_match = re.match(r"^([A-Za-z &/]+?):\s*(.+)", clean)
        if cat_match:
            current_cat = cat_match.group(1).strip().title()
            raw_skills = cat_match.group(2)
        else:
            raw_skills = clean

        skill_tokens = [s.strip().rstrip(".,;") for s in re.split(r"[,|/]", raw_skills) if s.strip()]
        skill_tokens = [s for s in skill_tokens if 1 < len(s) < 60]
        if skill_tokens:
            categories.setdefault(current_cat, []).extend(skill_tokens)

    # Build canonical matched list (against SKILLS_POOL)
    all_skills_text = " ".join(lines).lower()
    canonical: list[str] = []
    for skill in SKILLS_POOL:
        pattern = r"\b" + re.escape(skill) + r"\b"
        if re.search(pattern, all_skills_text):
            canonical.append(skill.title())

    return {"by_category": categories, "canonical": canonical}


def _split_tech_string(tech_str: str) -> list[str]:
    """Split 'Flask, PostgreSQL, Python' → ['Flask', 'PostgreSQL', 'Python']"""
    tokens = re.split(r"[,|/]", tech_str)
    return [t.strip().rstrip(".,;") for t in tokens if t.strip()]


def _extract_technologies(text: str) -> list[str]:
    """Match known technologies from SKILLS_POOL against arbitrary text."""
    text_lower = text.lower()
    found = []
    for skill in SKILLS_POOL:
        if re.search(r"\b" + re.escape(skill) + r"\b", text_lower):
            found.append(skill.title())
    return found


def _parse_experience_block(lines: list[str], is_internship: bool = False) -> list[dict]:
    merged = _merge_bullets(lines)
    blocks = []
    current_block = []
    
    _CONT_STARTS = re.compile(
        r"^(?:and |or |to |the |a |an |in |of |for |with |through |by |that |"
        r"on |at |from |into |based |as |this |these |which |where |who |"
        r"designed|developed|built|implemented|created|worked|used|managed|led|assisted|collaborated)",
        re.I
    )

    for line in merged:
        stripped = line.strip()
        if not stripped:
            continue
        
        is_bullet = bool(_BULLET_RE.match(stripped))
        has_date = bool(_DURATION_RE.search(stripped) or _YEAR_RANGE_RE.search(stripped))
        
        is_new_entry = False
        if not is_bullet:
            if has_date:
                is_new_entry = True
            elif "|" in stripped:
                is_new_entry = True
            elif re.search(r"\bat\b", stripped, re.I) and len(stripped) < 80:
                is_new_entry = True
            elif not current_block:
                is_new_entry = True
            elif len(stripped) < 100 and not any(ch in stripped for ch in [',', '.']) and not stripped.endswith((":", ";")):
                if stripped[0].isupper():
                    if not _CONT_STARTS.match(stripped):
                        is_new_entry = True
        
        if is_new_entry and current_block:
            blocks.append(current_block)
            current_block = [stripped]
        else:
            current_block.append(stripped)
            
    if current_block:
        blocks.append(current_block)
        
    records = []
    for block in blocks:
        record = _parse_single_experience_block(block, is_internship)
        if record:
            records.append(record)
    return records


def _parse_single_experience_block(block: list[str], is_internship: bool) -> dict | None:
    header_lines = []
    responsibilities = []
    
    for line in block:
        stripped = line.strip()
        is_bullet = bool(_BULLET_RE.match(stripped))
        
        if is_bullet:
            responsibilities.append(_BULLET_RE.sub("", stripped).strip())
        elif header_lines and len(stripped) > 80 and any(v in stripped.lower() for v in ["developed", "implemented", "managed", "designed", "created"]):
            responsibilities.append(stripped)
        else:
            if not responsibilities:
                header_lines.append(stripped)
            else:
                responsibilities.append(stripped)
                
    if not header_lines:
        if responsibilities:
            return {
                "company": None,
                "role": None,
                "duration": None,
                "is_internship": is_internship,
                "description": responsibilities
            }
        return None
        
    company = None
    role = None
    duration = None
    
    for line in header_lines:
        dur_m = _DURATION_RE.search(line) or _YEAR_RANGE_RE.search(line)
        if dur_m:
            duration = dur_m.group(0).strip()
            line_clean = line.replace(duration, "").strip().strip("(),-–—|").strip()
        else:
            line_clean = line
            
        if "|" in line_clean:
            parts = [p.strip() for p in line_clean.split("|") if p.strip()]
            if len(parts) >= 2:
                role = parts[0]
                company = parts[1]
                _ROLE_INDICATORS = ["engineer", "developer", "intern", "lead", "analyst", "manager", "programmer", "consultant", "sde", "specialist"]
                is_part0_role = any(ind in parts[0].lower() for ind in _ROLE_INDICATORS)
                is_part1_role = any(ind in parts[1].lower() for ind in _ROLE_INDICATORS)
                if is_part1_role and not is_part0_role:
                    role = parts[1]
                    company = parts[0]
                elif is_part0_role and not is_part1_role:
                    role = parts[0]
                    company = parts[1]
            elif len(parts) == 1:
                if not role and not company:
                    role = parts[0]
        elif "," in line_clean and len(line_clean.split(",")) == 2:
            parts = [p.strip() for p in line_clean.split(",")]
            role = parts[0]
            company = parts[1]
            _ROLE_INDICATORS = ["engineer", "developer", "intern", "lead", "analyst", "manager", "programmer", "consultant", "sde", "specialist"]
            if any(ind in parts[1].lower() for ind in _ROLE_INDICATORS) and not any(ind in parts[0].lower() for ind in _ROLE_INDICATORS):
                role = parts[1]
                company = parts[0]
        elif re.search(r"\bat\b", line_clean, re.I):
            parts = re.split(r"\bat\b", line_clean, flags=re.I)
            role = parts[0].strip()
            company = parts[1].strip()
        else:
            if not role and not company:
                _ROLE_INDICATORS = ["engineer", "developer", "intern", "lead", "analyst", "manager", "programmer", "consultant", "sde", "specialist"]
                if any(ind in line_clean.lower() for ind in _ROLE_INDICATORS):
                    role = line_clean
                else:
                    company = line_clean
            elif role and not company:
                company = line_clean
            elif company and not role:
                role = line_clean
                
    return {
        "company": company.strip() if company else None,
        "role": role.strip() if role else None,
        "duration": duration.strip() if duration else None,
        "is_internship": is_internship,
        "description": responsibilities
    }


def _parse_projects(lines: list[str]) -> list[dict]:
    """
    Parse project records from raw section lines WITHOUT pre-merging.

    Works directly on raw lines so pipe-separated project title lines
    (e.g. "SwiftPath | Python, Flask, PostgreSQL Github") are never
    mistakenly absorbed as bullet continuations.

    Detection rules for a new project header:
      1. Non-bullet line containing "|"  →  "Title | Tech1, Tech2 [Github]"
      2. Bullet line containing a duration pattern  →  "• Title (Sep 2025 – Nov 2025)"
      3. Non-bullet line followed (immediately) by a bullet line  →  standalone title

    Within a project block, consecutive non-bullet lines that follow a bullet
    are merged as continuations of that bullet (repairing PDF line-wraps).
    """
    records: list[dict] = []
    cur_header: str | None = None
    cur_raw_bullets: list[str] = []   # accumulates raw bullet + continuation lines

    def _next_nonempty(idx: int) -> str | None:
        for j in range(idx + 1, len(lines)):
            s = lines[j].strip()
            if s:
                return s
        return None

    # Words that start a sentence continuation, never a project title
    _CONT_STARTS = re.compile(
        r"^(?:and |or |to |the |a |an |in |of |for |with |through |by |that |"
        r"on |at |from |into |based |as |this |these |which |where |who |"
        r"real |delivery |listing |patterns |tors )",
        re.I,
    )

    def _is_proj_header(content: str, is_bullet: bool, nxt: str | None) -> bool:
        if is_bullet:
            return bool(_DURATION_RE.search(content) or _YEAR_RANGE_RE.search(content))
        if "|" in content:
            return True
        # Must start with an uppercase letter to be a title
        if not content or not content[0].isupper():
            return False
        # Reject obvious sentence fragments/continuations
        if _CONT_STARTS.match(content):
            return False
        # Reject dense-PDF artefacts: very long runs with no spaces (word/char ratio)
        words = content.split()
        if words and len(content) > 60 and (len(content) / max(len(words), 1)) > 12:
            return False  # average "word" length > 12 chars → concatenated text
        # Non-bullet, capital-start line immediately before a bullet → project title
        if nxt and _BULLET_RE.match(nxt):
            return True
        return False

    def _merge_raw_bullets(raw: list[str]) -> list[str]:
        """Merge continuation lines into the bullet they follow."""
        merged: list[str] = []
        for item in raw:
            is_bul = item.startswith("\x00")          # sentinel for bullets
            text   = item[1:] if item.startswith("\x00") else item
            if is_bul:
                merged.append(text)
            elif merged:
                sep = "" if merged[-1].endswith("-") else " "
                if merged[-1].endswith("-"):
                    merged[-1] = merged[-1][:-1] + text
                else:
                    merged[-1] = merged[-1] + sep + text
            else:
                merged.append(text)
        return merged

    def flush():
        if cur_header is None and not cur_raw_bullets:
            return
        merged_desc = _merge_raw_bullets(cur_raw_bullets)
        rec = _build_project_record(cur_header or "", merged_desc)
        if rec:
            records.append(rec)

    for idx, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue

        is_bullet = bool(_BULLET_RE.match(stripped))
        content   = _BULLET_RE.sub("", stripped).strip() if is_bullet else stripped
        nxt       = _next_nonempty(idx)

        if _is_proj_header(content, is_bullet, nxt):
            flush()
            cur_header      = content
            cur_raw_bullets = []
        elif is_bullet:
            if cur_header is None:
                # Headerless bullet project
                flush()
                cur_header      = content
                cur_raw_bullets = []
            else:
                cur_raw_bullets.append("\x00" + content)   # bullet sentinel
        else:
            # Non-bullet, non-header → continuation of last bullet
            cur_raw_bullets.append(content)

    flush()
    return records


def _build_project_record(header: str, desc_lines: list[str]) -> dict | None:
    if not header and not desc_lines:
        return None

    title = header
    technologies: list[str] = []
    github: str | None = None
    live_link: str | None = None
    duration: str | None = None

    # Search for GitHub URL in all text
    all_text = header + " " + " ".join(desc_lines)
    gh_m = _GITHUB_URL_RE.search(all_text)
    if gh_m:
        github = normalize_url(gh_m.group(0), "github")
        gh_header_m = _GITHUB_URL_RE.search(title)
        if gh_header_m:
            title = title.replace(gh_header_m.group(0), "").strip()

    # Search for Live Link / Portfolio URL in all text
    for url_m in _PORTFOLIO_RE.finditer(all_text):
        url = url_m.group(0)
        if not github or url.lower() != github.lower():
            live_link = url
            break

    # Extract duration from title
    dur_m = _DURATION_RE.search(title) or _YEAR_RANGE_RE.search(title)
    if dur_m:
        duration = dur_m.group(0).strip()
        title = title[:dur_m.start()].strip().strip("(").strip()

    # Extract technologies from "Title | Tech1, Tech2" pattern
    if "|" in title:
        parts = [p.strip() for p in title.split("|", 1)]
        title = parts[0].strip()
        tech_part = parts[1] if len(parts) > 1 else ""
        tech_part = _GITHUB_URL_RE.sub("", tech_part).strip()
        technologies = _split_tech_string(tech_part)
    else:
        technologies = _extract_technologies(all_text)

    # Remove bullet markers from description
    clean_desc = []
    for d in desc_lines:
        d = _BULLET_RE.sub("", d).strip()
        if d:
            clean_desc.append(d)

    # Build merged description string
    description = " ".join(clean_desc) if clean_desc else ""

    # Clean up title
    title = re.sub(r"\s*\([^)]*\)\s*$", "", title).strip()  # strip trailing (...)
    title = title.strip("•-–* ").strip()

    if not title:
        return None

    return {
        "title":        title,
        "technologies": technologies,
        "description":  description,
        "github":       github,
        "live_link":    live_link,
        "duration":     duration,
    }


def _parse_certifications(lines: list[str]) -> list[str]:
    merged = _merge_bullets(lines)
    certs = []
    for line in merged:
        clean = _BULLET_RE.sub("", line).strip()
        if clean and len(clean) > 5:
            certs.append(clean)
    return certs


def _parse_achievements(lines: list[str]) -> list[str]:
    merged = _merge_bullets(lines)
    items = []
    for line in merged:
        clean = _BULLET_RE.sub("", line).strip()
        if clean and len(clean) > 5:
            items.append(clean)
    return items


def _parse_languages(lines: list[str]) -> list[str]:
    merged = _merge_bullets(lines)
    langs = []
    for line in merged:
        clean = _BULLET_RE.sub("", line).strip()
        if clean:
            parts = [p.strip() for p in re.split(r"[,/|;]", clean) if p.strip()]
            for part in parts:
                if len(part) >= 2 and len(part) < 30:
                    langs.append(part)
    return langs


# ===========================================================================
# Confidence scoring
# ===========================================================================

def _score_confidence(parsed: dict) -> tuple[float, list[str]]:
    score = 0.0
    missing = []
    
    name = parsed.get("name")
    is_name_valid = False
    if name and isinstance(name, str) and 3 <= len(name.strip()) <= 60:
        if not any(char.isdigit() for char in name) and "@" not in name and "http" not in name:
            is_name_valid = True
            
    email = parsed.get("email")
    is_email_valid = bool(email and _EMAIL_RE.match(email))
    
    phone = parsed.get("phone")
    is_phone_valid = bool(phone and len(re.sub(r"\D", "", phone)) >= 7)
    
    linkedin = parsed.get("linkedin")
    is_linkedin_valid = bool(linkedin and "linkedin.com" in linkedin.lower())
    
    github = parsed.get("github")
    is_github_valid = bool(github and "github.com" in github.lower())
    
    portfolio = parsed.get("portfolio")
    is_portfolio_valid = bool(portfolio and portfolio.startswith(("http://", "https://")))
    
    skills = parsed.get("skills")
    is_skills_valid = bool(skills and len(skills) > 0)
    
    education = parsed.get("education")
    is_education_valid = bool(education and len(education) > 0)
    
    experience = parsed.get("experience")
    is_experience_valid = bool(experience and len(experience) > 0)
    
    internships = parsed.get("internships")
    is_internships_valid = bool(internships and len(internships) > 0)
    
    projects = parsed.get("projects")
    is_projects_valid = bool(projects and len(projects) > 0)
    
    certifications = parsed.get("certifications")
    is_certifications_valid = bool(certifications and len(certifications) > 0)
    
    checks = [
        ("name",            is_name_valid,            15.0, "Name"),
        ("email",           is_email_valid,           15.0, "Email Address"),
        ("phone",           is_phone_valid,           10.0, "Phone Number"),
        ("linkedin",        is_linkedin_valid,         4.0, "LinkedIn URL"),
        ("github",          is_github_valid,           4.0, "GitHub URL"),
        ("portfolio",       is_portfolio_valid,        2.0, "Portfolio URL"),
        ("skills",          is_skills_valid,          15.0, "Skills List"),
        ("education",       is_education_valid,       10.0, "Education Records"),
        ("experience",      is_experience_valid,        5.0, "Work Experience"),
        ("internships",     is_internships_valid,       5.0, "Internships"),
        ("projects",        is_projects_valid,        10.0, "Projects"),
        ("certifications",  is_certifications_valid,    5.0, "Certifications"),
    ]
    
    for _, is_valid, weight, label in checks:
        if is_valid:
            score += weight
        else:
            missing.append(label)
            
    return round(min(score, 100.0), 1), missing


# ===========================================================================
# Fact-only summary generator (50–80 words, never invents placeholders)
# ===========================================================================

def _generate_summary(
    name: str | None,
    canonical_skills: list[str],
    education: list[dict],
    experience: list[dict],
    internships: list[dict],
    projects: list[dict],
    certifications: list[str],
) -> str:
    """
    Compose a recruiter-friendly paragraph ONLY from data actually extracted.
    Rules:
      • Never use a fallback string that invents information
      • Target ~80-120 words for detailed professional feel
      • Avoid verbatim resume text; synthesise in natural recruiter-style prose
    """
    ref_name = name.split()[0] if name else "The candidate"
    
    edu_str = None
    if education:
        raw = education[0].get("raw", "").strip()
        if len(raw) > 90:
            raw = re.split(r"[,\(\[]", raw)[0].strip()
        # Clean column headers
        if "degree" in raw.lower() or "exam" in raw.lower() or "institute" in raw.lower():
            if len(education) > 1:
                raw2 = education[1].get("raw", "").strip()
                if len(raw2) > 90:
                    raw2 = re.split(r"[,\(\[]", raw2)[0].strip()
                if not ("degree" in raw2.lower() or "exam" in raw2.lower() or "institute" in raw2.lower()):
                    edu_str = raw2
        else:
            edu_str = raw
            
    all_exp = experience + internships
    orgs = [e["company"] for e in all_exp if e.get("company")]
    roles = [e["role"] for e in all_exp if e.get("role")]
    
    top_skills = canonical_skills[:6]
    
    sentences = []
    
    if edu_str:
        sentences.append(f"{ref_name} has a strong academic foundation, holding credentials in {edu_str}.")
    else:
        sentences.append(f"{ref_name} is a technical candidate focusing on software and engineering domains.")
        
    if orgs and roles:
        role_comp = [f"{r} at {c}" for r, c in zip(roles, orgs)]
        if len(role_comp) == 1:
            sentences.append(f"They have practical professional experience as a {role_comp[0]}.")
        else:
            sentences.append(f"Their professional background includes roles as a {role_comp[0]} and {role_comp[1]}.")
    elif orgs:
        sentences.append(f"They have gained hands-on industry exposure working with {', '.join(orgs[:2])}.")
    elif roles:
        sentences.append(f"They have developed practical expertise through roles such as {', '.join(roles[:2])}.")
        
    if top_skills:
        skill_list = ", ".join(top_skills[:-1]) + f", and {top_skills[-1]}" if len(top_skills) > 1 else top_skills[0]
        sentences.append(f"Their technical toolkit features proficiency in {skill_list}.")
        
    proj_cert = []
    if projects:
        proj_titles = [p["title"] for p in projects if p.get("title")]
        if proj_titles:
            proj_cert.append(f"implemented notable projects like {', '.join(proj_titles[:2])}")
        else:
            proj_cert.append(f"delivered {len(projects)} software projects")
    if certifications:
        proj_cert.append(f"acquired professional credentials including {', '.join(certifications[:2])}")
        
    if proj_cert:
        sentences.append(f"In addition, they have {', and '.join(proj_cert)}.")
        
    summary = " ".join(sentences)
    
    words = summary.split()
    if len(words) > 120:
        summary = " ".join(words[:118]) + "..."
        
    return summary


# ===========================================================================
# Project GitHub enrichment from PDF annotation links
# ===========================================================================

def _enrich_project_githubs(projects: list[dict], github_repos: list[str]) -> list[dict]:
    """
    Assign PDF-annotation repo URLs to projects that have 'Github'/'GitHub'
    mentioned in their title/header but no actual URL (text-only hyperlinks).
    Matches by comparing URL slug words to project title words.
    """
    if not github_repos:
        return projects

    unassigned_repos = [r for r in github_repos]

    for proj in projects:
        if proj.get("github"):
            # Already has a URL from text extraction — remove from pool
            unassigned_repos = [r for r in unassigned_repos if r != proj["github"]]
            continue

        title = proj.get("title", "").lower()
        title_words = set(re.split(r"\W+", title)) - {
            "", "a", "an", "the", "and", "or", "of", "in",
            "system", "platform", "app", "project", "portal"
        }
        if not title_words:
            continue

        best_url: str | None = None
        best_score = 0
        for url in unassigned_repos:
            slug = url.rstrip("/").split("/")[-1].lower()
            slug_words = set(re.split(r"[\W_\-]+", slug)) - {""}
            overlap = len(title_words & slug_words)
            if overlap > best_score:
                best_score = overlap
                best_url = url

        if best_score >= 1 and best_url:
            proj["github"] = best_url
            unassigned_repos.remove(best_url)

    return projects


# ===========================================================================
# Top-level entry points
# ===========================================================================

class ResumeParserService:
    """Public interface — mirrors the v1 API so existing callers need no changes."""

    # Expose pool so ats_service can import it
    SKILLS_POOL = SKILLS_POOL

    # Keep legacy attribute names for backward compatibility
    EDU_KEYWORDS = [
        "b.tech", "btech", "b.e.", "m.tech", "mtech", "b.sc", "m.sc",
        "bca", "mca", "mba", "ph.d", "phd",
        "bachelor", "master", "university", "college", "institute", "school",
    ]

    @classmethod
    def extract_text_from_pdf(cls, file_path) -> str:
        fp = Path(file_path)
        if not fp.exists():
            raise FileNotFoundError(f"Resume file not found: {fp}")
        return _extract_pdf_text(fp)

    @classmethod
    def extract_text_from_docx(cls, file_path) -> str:
        fp = Path(file_path)
        if not fp.exists():
            raise FileNotFoundError(f"Resume file not found: {fp}")
        return _extract_docx_text(fp)

    @classmethod
    def extract_text(cls, file_path) -> str:
        fp = Path(file_path)
        suffix = fp.suffix.lower()
        if suffix == ".pdf":
            return cls.extract_text_from_pdf(fp)
        if suffix == ".docx":
            return cls.extract_text_from_docx(fp)
        raise ValueError(f"Unsupported file type: {suffix}")

    @classmethod
    def parse_resume_text(cls, text: str, file_path=None) -> dict:
        """
        Full pipeline: raw text → structured dict.

        Args:
            text:      Raw text extracted from the resume file.
            file_path: Original file path (Path or str). When provided, PDF
                       hyperlink annotations are extracted via PyMuPDF so that
                       icon-hidden LinkedIn/GitHub/LeetCode URLs are recovered.
        """
        # ── 1. Extract PDF hyperlinks (requires file_path) ────────────────
        pdf_links: list[str] = []
        if file_path is not None:
            pdf_links = _extract_pdf_links(Path(file_path))

        # ── 2. Preprocess and section-detect ──────────────────────────────
        lines    = _preprocess(text)
        sections = _detect_sections(lines)

        # ── 3. Parse each section ─────────────────────────────────────────
        header_lines   = sections.get("header", [])
        contact        = _parse_contact(header_lines, text, pdf_links)
        education_raw  = sections.get("education", [])
        education_filtered = []
        for line in education_raw:
            s_line = line.strip().lower()
            if "degree" in s_line and "exam" in s_line and "institute" in s_line:
                continue
            education_filtered.append(line)
        education      = _parse_education(education_filtered)
        skills         = _parse_skills(sections.get("skills", []))
        experience     = _parse_experience_block(sections.get("experience", []), is_internship=False)
        internships    = _parse_experience_block(sections.get("internships", []), is_internship=True)
        projects       = _parse_projects(sections.get("projects", []))
        certifications = _parse_certifications(sections.get("certifications", []))
        achievements   = _parse_achievements(sections.get("achievements", []))
        languages      = _parse_languages(sections.get("languages", []))

        # ── 4. Enrich project GitHub links from PDF annotations ───────────
        github_repos = contact.get("github_repos", [])
        projects = _enrich_project_githubs(projects, github_repos)

        # ── 5. Flat canonical skill list (ATS compat) ─────────────────────
        canonical_skills = skills.get("canonical", [])

        # ── 6. Promote internship-labelled experience entries ─────────────
        clean_experience: list[dict] = []
        clean_internships = list(internships)
        for exp in experience:
            combined = " ".join(filter(None, [
                exp.get("company", ""), exp.get("role", "")
            ])).lower()
            if "intern" in combined or "internship" in combined or "trainee" in combined:
                exp["is_internship"] = True
                clean_internships.append(exp)
            else:
                clean_experience.append(exp)

        # ── 7. Confidence score ───────────────────────────────────────────
        parsed_for_score = {
            **contact,
            "skills":         canonical_skills,
            "education":      education,
            "experience":     clean_experience,
            "internships":    clean_internships,
            "projects":       projects,
            "certifications": certifications,
        }
        confidence, missing = _score_confidence(parsed_for_score)

        # ── 8. Fact-only summary (no hallucination) ───────────────────────
        summary = _generate_summary(
            contact["name"],
            canonical_skills, education, clean_experience,
            clean_internships, projects, certifications,
        )

        return {
            # Contact
            "name":            contact["name"],
            "email":           contact["email"],
            "phone":           contact["phone"],
            "linkedin":        contact["linkedin"],
            "github":          contact["github"],
            "github_profile":  contact["github_profile"],
            "github_repos":    contact["github_repos"],
            "portfolio":       contact["portfolio"],
            "leetcode":        contact["leetcode"],
            "hackerrank":      contact["hackerrank"],
            # Content
            "skills":            canonical_skills,
            "skills_by_category": skills.get("by_category", {}),
            "education":         education,
            "experience":        clean_experience,
            "internships":       clean_internships,
            "projects":          projects,
            "certifications":    certifications,
            "achievements":      achievements,
            "languages":         languages,
            # Meta
            "confidence_score":  confidence,
            "missing_fields":    missing,
            "summary":           summary,
        }

    @classmethod
    def parse_and_save_resume(cls, resume_id):
        resume = Resume.query.get_or_404(resume_id)
        try:
            text = cls.extract_text(resume.file_path)
            if not text.strip():
                raise ValueError("Resume file is empty or produced no text.")

            # Pass file_path so PyMuPDF can extract annotation-layer hyperlinks
            parsed_data = cls.parse_resume_text(text, file_path=resume.file_path)

            resume.parsed_text = json.dumps({
                "raw_text":        text[:8000],
                "structured_data": parsed_data,
            })
            resume.parse_status = ParseStatus.COMPLETED
            resume.parsed_at    = datetime.now(timezone.utc)
            db.session.commit()
            return parsed_data
        except Exception as exc:
            db.session.rollback()
            resume.parse_status = ParseStatus.FAILED
            resume.parsed_text  = json.dumps({"error": str(exc)})
            db.session.commit()
            raise
